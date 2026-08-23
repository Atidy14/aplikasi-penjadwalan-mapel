import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_manual_book():
    print("[INFO] Membuat dokumen Manual Book dalam format DOCX...")
    doc = Document()

    # Setup Margin Dokumen (Standar 1 Inch)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Warna Utama
    PRIMARY_COLOR = RGBColor(16, 185, 129)    # Emerald
    DARK_COLOR = RGBColor(15, 23, 42)         # Slate 900
    MUTED_COLOR = RGBColor(100, 116, 139)     # Slate 500
    ACCENT_COLOR = RGBColor(79, 70, 229)      # Indigo

    # ==========================================
    # HALAMAN SAMPUL (COVER PAGE)
    # ==========================================
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(40)
    title_p.paragraph_format.space_after = Pt(10)
    
    run_badge = title_p.add_run("BUKU PANDUAN PENGGUNA & ADMINISTRATOR\n")
    run_badge.font.name = 'Arial'
    run_badge.font.size = Pt(13)
    run_badge.font.bold = True
    run_badge.font.color.rgb = PRIMARY_COLOR

    run_title = title_p.add_run("SIP-MAPEL v2.0 PRO\nSISTEM PENJADWALAN MATA PELAJARAN SEKOLAH TERPADU")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(22)
    run_title.font.bold = True
    run_title.font.color.rgb = DARK_COLOR

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_after = Pt(30)
    run_sub = sub_p.add_run("Panduan Operasional End-to-End: Pengelolaan Data Master, Kertas Kerja, Ketersediaan Waktu Guru, Algoritma Auto-Generator Bebas Bentrok, Papan Penjadwalan, Serah Terima Guru & Audit Trail")
    run_sub.font.name = 'Calibri'
    run_sub.font.size = Pt(11)
    run_sub.font.color.rgb = MUTED_COLOR

    # Garis Pembatas
    border_table = doc.add_table(rows=1, cols=1)
    border_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = border_table.cell(0, 0)
    set_cell_background(cell, "10B981")
    cell.width = Inches(6.5)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("")

    # Info Institusi Cover
    info_p = doc.add_paragraph()
    info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_p.paragraph_format.space_before = Pt(40)
    info_p.paragraph_format.space_after = Pt(5)

    r_inst = info_p.add_run("PENGEMBANG SISTEM & ADMINISTRASI KURIKULUM\n")
    r_inst.font.name = 'Arial'
    r_inst.font.bold = True
    r_inst.font.size = Pt(12)
    r_inst.font.color.rgb = DARK_COLOR

    r_year = info_p.add_run("Tahun Ajaran Aktif: 2026/2027 Ganjil\nEdisi Resmi — Siap Implementasi Lapangan\nTanggal Rilis: Agustus 2026")
    r_year.font.name = 'Calibri'
    r_year.font.size = Pt(10.5)
    r_year.font.color.rgb = MUTED_COLOR

    doc.add_page_break()

    # ==========================================
    # LEMBAR INFORMASI & DAFTAR ISI
    # ==========================================
    h_toc = doc.add_heading(level=1)
    h_toc_run = h_toc.add_run("DAFTAR ISI & STRUKTUR PANDUAN")
    h_toc_run.font.name = 'Arial'
    h_toc_run.font.color.rgb = DARK_COLOR

    toc_items = [
        ("BAB 1: PENGENALAN & ARSITEKTUR SISTEM", "Latar belakang, arsitektur Next.js & Prisma, keunggulan fitur."),
        ("BAB 2: ALUR KERJA STANDAR (END-TO-END WORKFLOW)", "Diagram alir 5 tahap penyusunan jadwal sekolah bebas bentrok."),
        ("BAB 3: PANDUAN BERANDA & PORTAL LAYANAN UTAMA", "Navigasi antarmuka portal terpadu ala portal publik BPJS."),
        ("BAB 4: PENGELOLAAN DATA MASTER", "Modul Guru, Mata Pelajaran, dan Struktur Rombel Kelas (X, XI, XII)."),
        ("BAB 5: PENGATURAN KETERSEDIAAN & REQUEST IZIN GURU", "Konfigurasi Teacher Constraints (kalender hari/jam libur guru)."),
        ("BAB 6: PENYUSUNAN KERTAS KERJA (TEACHING LOAD)", "Matriks penugasan guru, mapel, kelas, dan alokasi jam/minggu."),
        ("BAB 7: EKSEKUSI AUTO-GENERATOR JADWAL CERDAS", "Algoritma Greedy Heuristic, pencarian slot otomatis, dan laporan."),
        ("BAB 8: PAPAN PENJADWALAN KELAS & VALIDASI INTERAKTIF", "Grid jadwal Senin-Sabtu, jam istirahat, validasi bentrok, progress target."),
        ("BAB 9: PROSEDUR SERAH TERIMA GURU (TEACHER HANDOVER)", "Mutasi guru di tengah semester dengan soft-delete & tracking historis."),
        ("BAB 10: AUDIT LOG & PENGAWASAN AKTIVITAS SISTEM", "Pencatatan rekam jejak mutasi jadwal dan aktvitas auto-generator."),
        ("BAB 11: PANDUAN TEKNIS & TROUBLESHOOTING ADMINISTRATOR", "Tips pemeliharaan database, ganti tahun ajaran, dan solusi kendala.")
    ]

    t_toc = doc.add_table(rows=len(toc_items) + 1, cols=2)
    t_toc.alignment = WD_TABLE_ALIGNMENT.CENTER
    t_toc.autofit = False

    t_toc.cell(0, 0).paragraphs[0].add_run("Modul / Bab").font.bold = True
    t_toc.cell(0, 1).paragraphs[0].add_run("Deskripsi & Cakupan").font.bold = True
    set_cell_background(t_toc.cell(0, 0), "F1F5F9")
    set_cell_background(t_toc.cell(0, 1), "F1F5F9")
    t_toc.cell(0, 0).width = Inches(2.8)
    t_toc.cell(0, 1).width = Inches(3.7)

    for idx, (bab, desc) in enumerate(toc_items):
        row_cells = t_toc.rows[idx + 1].cells
        r1 = row_cells[0].paragraphs[0].add_run(bab)
        r1.font.bold = True
        r1.font.size = Pt(9.5)
        r2 = row_cells[1].paragraphs[0].add_run(desc)
        r2.font.size = Pt(9.5)
        set_cell_margins(row_cells[0], top=80, bottom=80, left=100, right=100)
        set_cell_margins(row_cells[1], top=80, bottom=80, left=100, right=100)

    doc.add_page_break()

    # Helper function untuk menambahkan section dengan format rapi
    def add_chapter_header(number, title):
        h = doc.add_heading(level=1)
        r = h.add_run(f"BAB {number}: {title.upper()}")
        r.font.name = 'Arial'
        r.font.size = Pt(14)
        r.font.bold = True
        r.font.color.rgb = DARK_COLOR
        h.paragraph_format.space_before = Pt(16)
        h.paragraph_format.space_after = Pt(8)

    def add_section_header(title):
        h = doc.add_heading(level=2)
        r = h.add_run(title)
        r.font.name = 'Arial'
        r.font.size = Pt(12)
        r.font.bold = True
        r.font.color.rgb = PRIMARY_COLOR
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(4)

    def add_callout(text, box_type="NOTE"):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        c = tbl.cell(0, 0)
        c.width = Inches(6.5)
        bg = "ECFDF5" if box_type == "SUCCESS" else ("FEF3C7" if box_type == "WARNING" else "F8FAFC")
        set_cell_background(c, bg)
        set_cell_margins(c, top=120, bottom=120, left=180, right=180)
        p = c.paragraphs[0]
        prefix = "💡 PETUNJUK: " if box_type == "NOTE" else ("⚠️ PERINGATAN: " if box_type == "WARNING" else "✅ REKOMENDASI: ")
        r_pre = p.add_run(prefix)
        r_pre.font.bold = True
        r_pre.font.size = Pt(9.5)
        r_txt = p.add_run(text)
        r_txt.font.size = Pt(9.5)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)

    def add_image_if_exists(img_name, caption):
        img_path = os.path.join('screenshots', img_name)
        if os.path.exists(img_path):
            doc.add_paragraph().paragraph_format.space_before = Pt(6)
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.paragraph_format.space_after = Pt(4)
            p_img.add_run().add_picture(img_path, width=Inches(6.2))
            
            p_cap = doc.add_paragraph()
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_cap.paragraph_format.space_after = Pt(12)
            r_cap = p_cap.add_run(f"Gambar: {caption}")
            r_cap.font.name = 'Calibri'
            r_cap.font.size = Pt(9)
            r_cap.font.italic = True
            r_cap.font.color.rgb = MUTED_COLOR

    # ==========================================
    # BAB 1: PENGENALAN & ARSITEKTUR
    # ==========================================
    add_chapter_header(1, "Pengenalan & Arsitektur Sistem")
    
    p = doc.add_paragraph("Sistem Informasi Penjadwalan Mata Pelajaran (SIP-MAPEL v2.0 Pro) adalah platform perangkat lunak manajemen kurikulum sekolah terpadu yang dirancang khusus untuk memecahkan kompleksitas penyusunan jadwal belajar mengajar secara otomatis, akurat, dan bebas bentrok.")
    p.paragraph_format.space_after = Pt(6)

    add_section_header("1.1 Arsitektur & Teknologi Pilihan")
    p = doc.add_paragraph("Aplikasi ini dibangun menggunakan arsitektur teknologi modern berstandar enterprise:")
    
    tech_items = [
        ("Next.js (App Router)", "Framework React generasi terbaru dengan Server Actions untuk eksekusi mutasi backend instan tanpa jeda rendering."),
        ("TypeScript (Strict Type)", "Menjamin keamanan tipe data di seluruh komponen frontend dan logika bisnis backend."),
        ("Tailwind CSS & Shadcn UI", "Antarmuka responsif, modern, dan modular dengan navigasi portal terpadu."),
        ("Prisma ORM (v7.9)", "Manajemen basis data relasional dengan skema terstruktur, driver adapter resmi, dan keamanan transaksi."),
        ("Timetabling Heuristic Algorithm", "Algoritma Greedy cerdas yang memetakan ratusan slot jadwal dengan mempertimbangkan beban guru, ketersediaan waktu, dan batas maksimal jam harian.")
    ]
    for name, desc in tech_items:
        p_item = doc.add_paragraph(style='List Bullet')
        p_item.paragraph_format.space_after = Pt(3)
        r1 = p_item.add_run(f"{name}: ")
        r1.font.bold = True
        p_item.add_run(desc)

    add_section_header("1.2 Keunggulan Utama")
    bullet_features = [
        "Satu Klik Auto-Generate: Menjadwalkan belasan kelas sekaligus dalam waktu 1-2 detik.",
        "Anti-Bentrok Mutlak: Mencegah seorang guru terjadwal di dua kelas berbeda pada hari dan jam yang sama.",
        "Manajemen Multi-Tahun Ajaran: Memisahkan arsip jadwal per tahun ajaran tanpa risiko data tertimpa.",
        "Teacher Availability Constraints: Mengakomodasi permintaan hari/jam libur atau izin mengajar khusus bagi guru.",
        "Soft-Turnover Handover: Serah terima jadwal guru di tengah semester dengan rekam jejak historis utuh.",
        "Visual Target Progress: Progress bar berwarna yang memantau ketercapaian target jam mengajar per mata pelajaran."
    ]
    for bf in bullet_features:
        p_bf = doc.add_paragraph(style='List Bullet')
        p_bf.paragraph_format.space_after = Pt(3)
        p_bf.add_run(bf)

    # ==========================================
    # BAB 2: ALUR KERJA STANDAR (END-TO-END)
    # ==========================================
    add_chapter_header(2, "Alur Kerja Standar (End-to-End Workflow)")
    
    p = doc.add_paragraph("Untuk memastikan jadwal tersusun sempurna tanpa kendala, ikuti 5 tahapan alur kerja standar berikut:")
    
    steps = [
        ("Tahap 1: Setup Data Master", "Input daftar Guru, Mata Pelajaran (beserta target jam/minggu), dan Struktur Rombel Kelas (X, XI, XII)."),
        ("Tahap 2: Pengaturan Izin Guru", "Atur jadwal ketersediaan waktu guru (Constraint) bagi guru yang memiliki jadwal kuliah, libur, atau tugas luar."),
        ("Tahap 3: Penyusunan Kertas Kerja", "Pasangkan guru pengajar dengan mata pelajaran dan kelas terkait di menu Kertas Kerja (Teaching Load)."),
        ("Tahap 4: Eksekusi Auto-Generator", "Jalankan fungsi Auto-Generate untuk menyusun seluruh blok jadwal secara otomatis dalam hitungan detik."),
        ("Tahap 5: Evaluasi & Penyesuaian", "Buka Papan Penjadwalan per kelas untuk meninjau distribusi jadwal, memeriksa progress target jam, atau melakukan penyesuaian manual.")
    ]

    t_step = doc.add_table(rows=len(steps) + 1, cols=2)
    t_step.alignment = WD_TABLE_ALIGNMENT.CENTER
    t_step.cell(0, 0).paragraphs[0].add_run("Tahapan").font.bold = True
    t_step.cell(0, 1).paragraphs[0].add_run("Aktivitas & Output").font.bold = True
    set_cell_background(t_step.cell(0, 0), "10B981")
    set_cell_background(t_step.cell(0, 1), "10B981")
    t_step.cell(0, 0).paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    t_step.cell(0, 1).paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    t_step.cell(0, 0).width = Inches(2.2)
    t_step.cell(0, 1).width = Inches(4.3)

    for idx, (st, desc) in enumerate(steps):
        rc = t_step.rows[idx + 1].cells
        rc[0].paragraphs[0].add_run(st).font.bold = True
        rc[1].paragraphs[0].add_run(desc)
        set_cell_margins(rc[0], top=80, bottom=80, left=100, right=100)
        set_cell_margins(rc[1], top=80, bottom=80, left=100, right=100)

    # ==========================================
    # BAB 3: PORTAL BERANDA TERPADU
    # ==========================================
    add_chapter_header(3, "Panduan Beranda & Portal Layanan Utama")
    p = doc.add_paragraph("Halaman Beranda (`/`) dirancang sebagai pintu gerbang terpadu (Landing Portal) dengan tata letak modern ala portal publik institusional BPJS Ketenagakerjaan. Seluruh informasi penting dan akses modul terangkum dalam satu tampilan terpusat.")
    
    add_image_if_exists("01_landing_page.png", "Tampilan Portal Beranda Terpadu SIP-MAPEL v2.0 Pro")

    add_section_header("3.1 Komponen Utama Beranda")
    features_home = [
        ("Status Bar Informasi", "Menampilkan status keaktifan sistem dan label Tahun Ajaran yang sedang berlaku (misal: 2026/2027 Ganjil)."),
        ("Hero Banner Aksi Cepat", "Menyediakan tombol pintas langsung menuju eksekusi 'Auto-Generate Jadwal' dan 'Kertas Kerja'."),
        ("Kartu Statistik Real-time", "Menampilkan metrik dinamis: Total Guru, Total Kelas/Rombel, Total Mapel, dan Jumlah Blok Jadwal yang telah terbit."),
        ("Grid Menu Layanan (8 Cards)", "Pintu masuk modular menuju seluruh fitur utama aplikasi dengan indikator visual warna dan badge status.")
    ]
    for title, desc in features_home:
        p_fh = doc.add_paragraph(style='List Bullet')
        p_fh.paragraph_format.space_after = Pt(3)
        r_fh = p_fh.add_run(f"{title}: ")
        r_fh.font.bold = True
        p_fh.add_run(desc)

    # ==========================================
    # BAB 4: PENGELOLAAN DATA MASTER
    # ==========================================
    add_chapter_header(4, "Pengelolaan Data Master")
    p = doc.add_paragraph("Data Master merupakan pondasi awal sebelum sistem dapat menyusun jadwal. Data Master terdiri dari Manajemen Guru, Mata Pelajaran, dan Struktur Rombel Kelas.")

    add_section_header("4.1 Manajemen Data Guru")
    p = doc.add_paragraph("Diakses melalui menu 'Guru' (`/master/teachers`). Administrator dapat menambahkan nama guru, mengubah status aktif/nonaktif, mengatur jam libur khusus, serta melakukan serah terima jadwal mengajar.")
    add_image_if_exists("02_data_guru.png", "Halaman Manajemen Data Guru dan Status")

    add_callout("Guru yang dinonaktifkan (Status: Nonaktif) tidak akan muncul pada pilihan pembuatan Kertas Kerja baru, namun rekam jejak historis mengajarnya tetap tersimpan aman di database.", "NOTE")

    add_section_header("4.2 Manajemen Mata Pelajaran")
    p = doc.add_paragraph("Diakses melalui menu 'Mata Pelajaran' (`/master/subjects`). Setiap mata pelajaran wajib memiliki parameter 'Target Jam Pelajaran per Minggu' (misal: Matematika 4 jam, Fisika 3 jam, dsb). Nilai ini akan menjadi acuan utama bagi validator dan algoritma penjadwalan.")
    add_image_if_exists("05_data_mapel.png", "Halaman Manajemen Mata Pelajaran & Target Jam")

    add_section_header("4.3 Manajemen Struktur Kelas (Rombel)")
    p = doc.add_paragraph("Diakses melalui menu 'Kelas' (`/master/classes`). Administrator dapat mendaftarkan rombongan belajar (misal: Kelas X-1 s.d. X-3, Kelas XI-1 s.d. XI-4, Kelas XII-1 s.d. XII-6). Pada tabel kelas, tersedia tombol hijau 'Atur Jadwal' untuk langsung membuka Papan Penjadwalan masing-masing kelas.")
    add_image_if_exists("06_data_kelas.png", "Halaman Manajemen Kelas & Tombol Akses Papan Jadwal")

    # ==========================================
    # BAB 5: PENGATURAN KETERSEDIAAN GURU
    # ==========================================
    add_chapter_header(5, "Pengaturan Ketersediaan & Request Izin Guru")
    p = doc.add_paragraph("Fitur 'Teacher Constraints' memungkinkan administrator memblokir hari atau jam tertentu di mana seorang guru berhalangan mengajar (misal: studi lanjut, tugas luar, atau jadwal piket).")

    add_image_if_exists("03_modal_izin_guru.png", "Modal Grid Kalender Pengaturan Ketersediaan Waktu Guru")

    add_section_header("5.1 Prosedur Pengaturan:")
    steps_constr = [
        "1. Buka menu Manajemen Guru (`/master/teachers`).",
        "2. Cari nama guru yang bersangkutan pada tabel, lalu klik tombol kuning 'Atur Libur/Izin'.",
        "3. Sebuah modal kalender interaktif (Senin s.d. Sabtu, Jam 1 s.d. 8) akan muncul di layar.",
        "4. Klik pada kotak jam yang ingin diblokir sehingga berubah menjadi warna MERAH bertanda silang [X].",
        "5. Klik kembali untuk mengizinkan (berubah menjadi HIJAU bertanda centang [✓]).",
        "6. Tutup modal. Perubahan tersimpan otomatis secara real-time!"
    ]
    for sc in steps_constr:
        p_sc = doc.add_paragraph()
        p_sc.paragraph_format.space_after = Pt(2)
        p_sc.add_run(sc)

    add_callout("Algoritma Auto-Generator secara otomatis membaca daftar larangan ini dan tidak akan pernah menempatkan jadwal guru tersebut di jam yang telah diblokir merah.", "SUCCESS")

    # ==========================================
    # BAB 6: PENYUSUNAN KERTAS KERJA
    # ==========================================
    add_chapter_header(6, "Penyusunan Kertas Kerja (Teaching Load)")
    p = doc.add_paragraph("Kertas Kerja (Teaching Load) di menu `/master/teaching-load` adalah blueprint utama yang mendefinisikan penugasan mengajar: 'Guru Siapa, Mengajar Mapel Apa, di Kelas Mana, dan Berapa Jam/Minggu'.")

    add_image_if_exists("07_kertas_kerja.png", "Halaman Kertas Kerja Penugasan Mengajar")

    add_section_header("6.1 Proteksi & Keunikan Data (Upsert Mechanism)")
    p = doc.add_paragraph("Sistem dilengkapi mekanisme proteksi ganda:")
    p_p1 = doc.add_paragraph(style='List Bullet')
    p_p1.add_run("Satu kombinasi Kelas dan Mapel hanya dapat dipegang oleh satu guru utama. Jika terjadi pergantian, sistem otomatis memperbarui (update) tanpa membuat data duplikat.")
    p_p2 = doc.add_paragraph(style='List Bullet')
    p_p2.add_run("Tersedia tombol hapus penugasan pada setiap baris jika ada alokasi beban mengajar yang dibatalkan.")

    # ==========================================
    # BAB 7: EKSEKUSI AUTO-GENERATOR JADWAL
    # ==========================================
    add_chapter_header(7, "Eksekusi Auto-Generator Jadwal Cerdas")
    p = doc.add_paragraph("Menu Auto-Generator (`/master/auto-generate`) adalah fitur pamungkas untuk menyusun seluruh jadwal sekolah dalam satu klik menggunakan pendekatan Algoritma Greedy Heuristic.")

    add_image_if_exists("08_auto_generator.png", "Halaman Eksekusi Auto-Generator dan Laporan Hasil")

    add_section_header("7.1 Mekanisme Kerja Algoritma")
    algo_steps = [
        ("Prioritas Beban Terberat", "Mapel dengan jam terbanyak (misal 4 jam) dijadwalkan terlebih dahulu untuk mempermudah distribusi ruang."),
        ("Pengecekan Bentrok Multi-Dimensi", "Memastikan Guru tidak bentrok di kelas lain dan Kelas tidak memiliki jadwal ganda pada jam yang sama."),
        ("Kepatuhan Constraint Guru", "Melewati seluruh slot jam yang telah ditandai merah/izin oleh guru bersangkutan."),
        ("Distribusi Harian Merata", "Membatasi maksimal 2 jam pelajaran untuk mapel yang sama dalam satu hari agar siswa tidak jenuh."),
        ("Penanganan Istirahat Otomatis", "Menghormati jeda istirahat di antara jam ke-4 dan jam ke-5.")
    ]
    for title, desc in algo_steps:
        p_as = doc.add_paragraph(style='List Bullet')
        p_as.paragraph_format.space_after = Pt(3)
        r_as = p_as.add_run(f"{title}: ")
        r_as.font.bold = True
        p_as.add_run(desc)

    # ==========================================
    # BAB 8: PAPAN PENJADWALAN KELAS
    # ==========================================
    add_chapter_header(8, "Papan Penjadwalan Kelas & Validasi Interaktif")
    p = doc.add_paragraph("Papan Penjadwalan (`/master/classes/[classId]/scheduler`) menampilkan matriks jadwal mingguan (Senin s.d. Sabtu, Jam 1 s.d. 8) lengkap dengan baris istirahat visual dan pemantau target jam.")

    add_image_if_exists("09_papan_penjadwalan_grid.png", "Papan Penjadwalan Grid dengan Progress Bar Pemenuhan Jam")

    add_section_header("8.1 Fitur Utama Papan Penjadwalan:")
    sched_feats = [
        ("Progress Bar Pemenuhan Jam", "Badge berwarna: HIJAU jika jam pas sesuai target, KUNING jika kurang jam, dan MERAH jika kelebihan jam."),
        ("Interaksi Klik Slot", "Klik pada sel kosong atau terisi untuk membuka panel editor bawah guna memilih Guru & Mapel atau Mengosongkan slot."),
        ("Validasi Bentrok Real-Time", "Jika admin menempatkan guru yang sudah mengajar di kelas lain pada jam yang sama, sistem langsung memunculkan pesan error merah dan membatalkan penyimpanan."),
        ("Baris Istirahat Visual", "Pemisah visual abu-abu di antara jam ke-4 dan jam ke-5 untuk memudahkan pembacaan jadwal.")
    ]
    for title, desc in sched_feats:
        p_sf = doc.add_paragraph(style='List Bullet')
        p_sf.paragraph_format.space_after = Pt(3)
        r_sf = p_sf.add_run(f"{title}: ")
        r_sf.font.bold = True
        p_sf.add_run(desc)

    # ==========================================
    # BAB 9: PROSEDUR SERAH TERIMA GURU (HANDOVER)
    # ==========================================
    add_chapter_header(9, "Prosedur Serah Terima Guru (Teacher Handover)")
    p = doc.add_paragraph("Fitur 'Serah Terima Guru' menangani fenomena turnover pengajar di tengah semester (misal: mutasi, pensiun, atau cuti panjang) tanpa menghapus data historis.")

    add_image_if_exists("04_modal_serah_terima.png", "Modal Serah Terima Jadwal Guru Antar-Pengajar")

    add_section_header("9.1 Cara Melakukan Serah Terima:")
    steps_ho = [
        "1. Buka menu Manajemen Guru (`/master/teachers`).",
        "2. Klik tombol ungu 'Serah Terima Guru (Handover)' di kanan atas.",
        "3. Pilih 'Guru Lama' yang akan digantikan dan 'Guru Baru' pengganti.",
        "4. Tentukan 'Tanggal Berlaku' (Effective Date).",
        "5. Klik 'Proses Handover'."
    ]
    for sho in steps_ho:
        p_sho = doc.add_paragraph()
        p_sho.paragraph_format.space_after = Pt(2)
        p_sho.add_run(sho)

    add_callout("Sistem secara otomatis menutup jadwal guru lama dengan menyetel `validUntil` = (Tanggal Berlaku - 1 hari), lalu membuat jadwal baru untuk guru pengganti dengan `validFrom` = Tanggal Berlaku. Seluruh histori tercatat di Audit Log!", "NOTE")

    # ==========================================
    # BAB 10: AUDIT LOG & REKAM JEJAK
    # ==========================================
    add_chapter_header(10, "Audit Log & Pengawasan Aktivitas Sistem")
    p = doc.add_paragraph("Halaman Audit Log (`/master/audit`) mencatat seluruh aktivitas krusial sistem seperti proses auto-generate jadwal, mutasi guru (handover), dan modifikasi slot.")

    add_image_if_exists("10_audit_log.png", "Halaman Riwayat Perubahan dan Audit Trail")

    p_log = doc.add_paragraph("Setiap baris log menampilkan: Jenis Tindakan (Action Badge), Pelaksana (Performed By), Tanggal & Jam Kejadian (WIB), serta rincian detail peristiwa dalam bahasa yang mudah dipahami.")

    # ==========================================
    # BAB 11: PANDUAN TEKNIS & TROUBLESHOOTING
    # ==========================================
    add_chapter_header(11, "Panduan Teknis & Troubleshooting Administrator")
    
    faq_items = [
        ("Bagaimana cara mengganti ke Tahun Ajaran Baru?", "Di database tabel `AcademicYear`, buat record baru (misal: '2027/2028 Ganjil') dan set `isActive: true`. Buat kelas-kelas baru untuk tahun tersebut, susun Kertas Kerja, lalu jalankan Auto-Generator. Jadwal tahun lama tetap aman tersimpan."),
        ("Apa yang harus dilakukan jika ada slot gagal terjadwal (Unassigned Slot)?", "Buka menu Papan Penjadwalan kelas terkait. Periksa apakah guru bersangkutan memiliki jam izin yang terlalu padat. Masukkan jadwal secara manual ke slot kosong yang disepakati."),
        ("Bagaimana jika tidak sengaja menekan tombol Auto-Generator?", "Tekan tombol merah muda 'Undo / Pulihkan Jadwal Sebelumnya' yang otomatis muncul di halaman Auto-Generator. Sistem akan memulihkan 100% kondisi jadwal lama dari snapshot backup otomatis."),
        ("Bagaimana cara menghubungkan aplikasi ke PostgreSQL untuk Production?", "Ubah konfigurasi `datasource db { provider = 'postgresql' }` di `schema.prisma`, ganti `DATABASE_URL` di `.env` dengan connection string PostgreSQL Anda, lalu jalankan `npx prisma db push`.")
    ]

    for q, a in faq_items:
        p_q = doc.add_paragraph()
        p_q.paragraph_format.space_before = Pt(8)
        p_q.paragraph_format.space_after = Pt(2)
        r_q = p_q.add_run(f"Q: {q}")
        r_q.font.bold = True
        r_q.font.color.rgb = DARK_COLOR
        
        p_a = doc.add_paragraph()
        p_a.paragraph_format.space_after = Pt(6)
        r_a = p_a.add_run(f"Jawab: {a}")
        r_a.font.color.rgb = RGBColor(51, 65, 85)

    # ==========================================
    # BAB 12: PENGUJIAN KEPATUHAN AUDITOR CISA
    # ==========================================
    add_chapter_header(12, "Pengujian Otomatis & Kepatuhan Standar Auditor CISA")
    p = doc.add_paragraph("Sistem SIP-MAPEL v2.0 Pro telah dilengkapi suite pengujian otomatis (Automated Unit Testing) yang dirancang sesuai standar ISACA ITAF (IT Assurance Framework) dan ISO/IEC 25010 Software Quality Model.")

    add_section_header("12.1 Kontrol Audit yang Diverifikasi (100% PASS):")
    cisa_controls = [
        ("AC-1.1 Data Integrity (Teacher Multi-Booking)", "Memastikan zero-conflict: tidak ada guru mengajar di 2 kelas pada jam yang sama."),
        ("AC-1.2 Data Integrity (Class Collision)", "Memastikan zero-collision: tidak ada kelas memiliki 2 mapel bersamaan di ruangan yang sama."),
        ("AC-2.1 Constraint Compliance (Teacher Availability)", "Memverifikasi kepatuhan mutlak terhadap batasan libur/izin guru."),
        ("AC-2.2 Pedagogical Fatigue Limit", "Memastikan batasan maksimal 2 jam pelajaran untuk mapel yang sama per hari."),
        ("AC-3.1 Traceability & Non-Repudiation", "Memastikan pencatatan audit log otomatis lengkap dengan aktor dan timestamp WIB."),
        ("AC-4.1 State Recovery & Rollback", "Menguji keakuratan pemulihan snapshot jadwal (Undo) dengan presisi delta = 0."),
        ("AC-4.2 Non-Destructive Turnover", "Memverifikasi kontinuitas historis masa berlaku jadwal guru (validFrom/validUntil)."),
        ("AC-5.1 Compound Constraint & Idempotency", "Mencegah duplikasi data ganda pada matriks Kertas Kerja.")
    ]
    for c_title, c_desc in cisa_controls:
        p_c = doc.add_paragraph(style='List Bullet')
        p_c.paragraph_format.space_after = Pt(3)
        r_c = p_c.add_run(f"{c_title}: ")
        r_c.font.bold = True
        p_c.add_run(c_desc)

    add_callout("Administrator atau Auditor Eksternal dapat menjalankan verifikasi sistem sewaktu-waktu dengan menjalankan perintah: `npm test` di terminal.", "SUCCESS")

    # Simpan File Dokumen DOCX
    output_filename = "MANUAL_BOOK_SIP_MAPEL_v2.1_ENTERPRISE.docx"
    doc.save(output_filename)
    print(f"[SUCCESS] Dokumen Manual Book berhasil dibuat: {output_filename}")

if __name__ == "__main__":
    create_manual_book()
